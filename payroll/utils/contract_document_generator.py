from io import BytesIO
from datetime import datetime, timezone
from pathlib import Path
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.units import cm
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import logging

logger = logging.getLogger(__name__)


class ContractDocumentGenerator:
    """Generador de documentos PDF y DOCX para contratos establecidos."""
    
    @staticmethod
    def _resolve_logo_path():
        """Resuelve la ruta del logo institucional."""
        try:
            current_dir = Path(__file__).resolve().parent
            payroll_dir = current_dir.parent
            project_root = payroll_dir.parent
            candidates = [
                project_root / 'core' / 'assets' / 'logo.jpg',
                project_root / 'core' / 'assets' / 'logo.png',
                payroll_dir / 'assets' / 'logo.jpg',
                payroll_dir / 'assets' / 'logo.png',
            ]
            for candidate in candidates:
                if candidate.exists():
                    return str(candidate)
        except Exception as e:
            logger.warning(f"Error al resolver logo: {str(e)}")
        return None
    
    @staticmethod
    def _format_user_name(user):
        """Formatea el nombre del usuario para mostrar."""
        if not user:
            return "Sistema"
        
        if isinstance(user, str):
            return user
        
        if isinstance(user, dict):
            name = (user.get('name') or user.get('first_name') or '').strip()
            fln = (user.get('first_last_name') or user.get('last_name') or '').strip()
            sln = (user.get('second_last_name') or '').strip()
            parts = [p for p in [name, fln, sln] if p]
            return ' '.join(parts) if parts else str(user)
        
        try:
            name = (getattr(user, 'name', None) or getattr(user, 'first_name', '') or '').strip()
            fln = (getattr(user, 'first_last_name', None) or getattr(user, 'last_name', '') or '').strip()
            sln = (getattr(user, 'second_last_name', '') or '').strip()
            parts = [p for p in [name, fln, sln] if p]
            return ' '.join(parts) if parts else str(user)
        except Exception:
            return str(user)
    
    @staticmethod
    def _format_payment_frequency_details(contract):
        """Formatea los detalles de frecuencia de pago según el tipo."""
        freq = contract.payment_frequency_type
        payments = list(contract.contract_payments.all())
        
        if freq == 'diario':
            return "Diario"
        elif freq == 'semanal':
            if payments and payments[0].id_day_of_week:
                day_name = payments[0].id_day_of_week.name
                return f"Semanal - Día: {day_name}"
            return "Semanal"
        elif freq == 'quincenal':
            dates = [str(p.date_payment) for p in payments if p.date_payment]
            if dates:
                return f"Quincenal - Días: {', '.join(dates)}"
            return "Quincenal"
        elif freq == 'mensual':
            if payments and payments[0].date_payment:
                return f"Mensual - Día: {payments[0].date_payment}"
            return "Mensual"
        return freq
    
    @staticmethod
    def _format_date(date_obj):
        """Formatea una fecha a string."""
        if not date_obj:
            return "N/A"
        if hasattr(date_obj, 'strftime'):
            return date_obj.strftime("%Y-%m-%d")
        return str(date_obj)
    
    @staticmethod
    def _format_datetime(dt_obj):
        """Formatea datetime a string."""
        if not dt_obj:
            return "N/A"
        if hasattr(dt_obj, 'strftime'):
            return dt_obj.strftime("%Y-%m-%d %H:%M:%S")
        return str(dt_obj)
    
    @staticmethod
    def generate_pdf(contract, downloader_user=None, logo_path=None):
        """
        Genera un PDF del contrato establecido.
        
        Args:
            contract: Instancia de EstablishedContract
            downloader_user: Usuario que descarga (str, dict u objeto)
            logo_path: Ruta al logo (opcional)
        
        Returns:
            bytes: Contenido del PDF
        """
        if logo_path is None:
            logo_path = ContractDocumentGenerator._resolve_logo_path()
        
        buffer = BytesIO()
        left_margin = 2 * cm
        right_margin = 2 * cm
        bottom_margin = 2 * cm
        
        # Calcular top margin considerando logo
        logo_height = 0
        header_margin = 1.5 * cm
        if logo_path:
            try:
                img_width = 3 * cm
                ir = ImageReader(logo_path)
                iw, ih = ir.getSize()
                ratio = ih / float(iw) if iw else 1.0
                logo_height = img_width * ratio
            except Exception:
                logo_height = 0
        
        top_margin = max(2 * cm, logo_height + header_margin + (1.0 * cm))
        
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=left_margin,
            rightMargin=right_margin,
            topMargin=top_margin,
            bottomMargin=bottom_margin,
        )
        
        styles = getSampleStyleSheet()
        small = ParagraphStyle('small', parent=styles['Normal'], fontSize=9, leading=11)
        table_cell_style = ParagraphStyle('table_cell', parent=styles['Normal'], fontSize=8, leading=10, wordWrap='LTR')
        
        
        # Contador de páginas para mostrar total (usando lista mutable)
        page_info = [0]  # [total_pages]
        
        # Encabezado con logo y título
        def header_footer(canvas, doc_content, total_pages=None):
            canvas.saveState()
            width, height = A4
            
            # Actualizar contador de páginas
            current_page = canvas.getPageNumber()
            if total_pages is None:
                # Primera pasada: contar páginas
                if current_page > page_info[0]:
                    page_info[0] = current_page
            
            # Logo en la parte superior derecha
            if logo_path:
                try:
                    img_width = 3 * cm
                    ir = ImageReader(logo_path)
                    iw, ih = ir.getSize()
                    ratio = ih / float(iw) if iw else 1.0
                    img_height = img_width * ratio
                    top_padding = 0.5 * cm
                    x = width - img_width - right_margin
                    y = height - img_height - top_padding
                    canvas.drawImage(logo_path, x, y, width=img_width, height=img_height, preserveAspectRatio=True, mask='auto')
                except Exception:
                    pass
            
            # Título del documento
            canvas.setFont('Helvetica-Bold', 16)
            title = f"Contrato de Trabajo {contract.contract_code}"
            title_width = canvas.stringWidth(title, 'Helvetica-Bold', 16)
            canvas.drawString((width - title_width) / 2, height - 1 * cm, title)
            
            # Fecha y hora de generación
            canvas.setFont('Helvetica', 9)
            gen_date = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S")
            date_width = canvas.stringWidth(f"Generado: {gen_date}", 'Helvetica', 9)
            canvas.drawString((width - date_width) / 2, height - 1.5 * cm, f"Generado: {gen_date}")
            
            # Pie de página con paginación completa
            page_num = current_page
            total = total_pages if total_pages is not None else page_info[0]
            user_name = ContractDocumentGenerator._format_user_name(downloader_user)
            
            # Mostrar paginación completa si tenemos el total
            if total > 0 and total_pages is not None:
                footer_text = f"Descargado por: {user_name} | Página {page_num} / {total}"
            else:
                footer_text = f"Descargado por: {user_name} | Página {page_num}"
            
            canvas.setFont('Helvetica', 8)
            footer_width = canvas.stringWidth(footer_text, 'Helvetica', 8)
            canvas.drawString((width - footer_width) / 2, bottom_margin / 2, footer_text)
            
            canvas.restoreState()
        
        # Callback para primera pasada (contar páginas)
        def on_page_first(canvas, doc_content):
            header_footer(canvas, doc_content, total_pages=None)
        
        # Callback para segunda pasada (con total conocido)
        def on_page_final(canvas, doc_content):
            header_footer(canvas, doc_content, total_pages=page_info[0])
        
        def create_story(available_width):
            story = []
            
            # Título principal
            story.append(Spacer(1, 0.5 * cm))
            story.append(Paragraph("CONTRATO DE TRABAJO", styles['Title']))
            story.append(Spacer(1, 0.3 * cm))
            
            # Sección 1: Generalidades del contrato
            story.append(Paragraph("1. GENERALIDADES DEL CONTRATO", styles['Heading2']))
            story.append(Spacer(1, 0.2 * cm))
            
            generalidades_data = [
                ["Código del contrato", contract.contract_code],
                ["Cargo", getattr(contract.id_employee_charge, 'name', 'N/A')],
                ["Descripción", contract.description or "N/A"],
                ["Tipo de contrato", getattr(contract.contract_type, 'name', 'N/A')],
                ["Fecha de inicio", ContractDocumentGenerator._format_date(contract.start_date)],
                ["Fecha de finalización", ContractDocumentGenerator._format_date(contract.end_date)],
                ["Frecuencia de pago", ContractDocumentGenerator._format_payment_frequency_details(contract)],
                ["Jornada laboral", getattr(contract.workday_type, 'name', 'N/A') if contract.workday_type else "N/A"],
                ["Modalidad de trabajo", getattr(contract.work_mode_type, 'name', 'N/A') if contract.work_mode_type else "N/A"],
                ["Estado", getattr(contract.established_contract_status, 'name', 'N/A')],
            ]
            
            generalidades_table = Table(generalidades_data, colWidths=[available_width * 0.4, available_width * 0.6], hAlign='LEFT')
            generalidades_table.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ]))
            story.append(generalidades_table)
            story.append(Spacer(1, 0.5 * cm))
            
            # Sección 2: Términos del contrato
            story.append(Paragraph("2. TÉRMINOS DEL CONTRATO", styles['Heading2']))
            story.append(Spacer(1, 0.2 * cm))
            
            currency_symbol = getattr(contract.currency_type, 'symbol', '') or ''
            salary_display = f"{currency_symbol}{contract.salary_base:,.2f}"
            
            terminos_data = [
                ["Tipo de salario", contract.salary_type],
                ["Salario base", salary_display],
                ["Moneda", getattr(contract.currency_type, 'name', 'N/A')],
                ["Período de prueba (días)", str(contract.trial_period_days) if contract.trial_period_days else "N/A"],
                ["Días de vacaciones", str(contract.vacation_days)],
                ["Vacaciones acumulativas", "Sí" if contract.cumulative_vacation else "No"],
                ["Fecha inicio acumulación", ContractDocumentGenerator._format_date(contract.start_cumulative_vacation) if contract.cumulative_vacation else "N/A"],
                ["Frecuencia de vacaciones (días)", str(contract.vacation_frequency_days) if contract.vacation_frequency_days else "N/A"],
                ["Días máximos de incapacidad", str(contract.maximum_disability_days)],
                ["Horas extras", f"{contract.overtime}" if contract.overtime else "0"],
                ["Período horas extras", contract.overtime_period or "N/A"],
                ["Período de notificación (días)", str(contract.notice_period_days) if contract.notice_period_days else "N/A"],
            ]
            
            if contract.minimum_hours:
                terminos_data.insert(3, ["Horas mínimas", str(contract.minimum_hours)])
            
            terminos_table = Table(terminos_data, colWidths=[available_width * 0.4, available_width * 0.6], hAlign='LEFT')
            terminos_table.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ]))
            story.append(terminos_table)
            story.append(Spacer(1, 0.5 * cm))
            
            # Sección 3: Deducciones
            deductions = list(contract.established_deductions.all())
            if deductions:
                story.append(Paragraph("3. DEDUCCIONES", styles['Heading2']))
                story.append(Spacer(1, 0.2 * cm))
                
                ded_header = [
                    Paragraph("Tipo", table_cell_style),
                    Paragraph("Tipo Monto", table_cell_style),
                    Paragraph("Valor", table_cell_style),
                    Paragraph("Aplicación", table_cell_style),
                    Paragraph("Inicio", table_cell_style),
                    Paragraph("Fin", table_cell_style),
                    Paragraph("Descripción", table_cell_style),
                    Paragraph("Cantidad", table_cell_style),
                ]
                
                ded_table_data = [ded_header]
                for ded in deductions:
                    ded_table_data.append([
                        Paragraph(getattr(ded.deduction_type, 'name', 'N/A'), table_cell_style),
                        Paragraph(ded.amount_type, table_cell_style),
                        Paragraph(f"{ded.amount_value}", table_cell_style),
                        Paragraph(ded.application_deduction_type, table_cell_style),
                        Paragraph(ContractDocumentGenerator._format_date(ded.start_date_deduction), table_cell_style),
                        Paragraph(ContractDocumentGenerator._format_date(ded.end_date_deductions), table_cell_style),
                        Paragraph(ded.description or "N/A", table_cell_style),
                        Paragraph(f"{ded.amount}" if ded.amount else "N/A", table_cell_style),
                    ])
                
                ded_col_widths = [available_width * 0.15, available_width * 0.12, available_width * 0.1, available_width * 0.12, 
                                 available_width * 0.1, available_width * 0.1, available_width * 0.16, available_width * 0.15]
                ded_table = Table(ded_table_data, colWidths=ded_col_widths, hAlign='LEFT', repeatRows=1)
                ded_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('LEFTPADDING', (0, 0), (-1, -1), 4),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                ]))
                story.append(ded_table)
                story.append(Spacer(1, 0.5 * cm))
            
            # Sección 4: Incrementos
            increases = list(contract.established_increases.all())
            if increases:
                story.append(Paragraph("4. INCREMENTOS", styles['Heading2']))
                story.append(Spacer(1, 0.2 * cm))
                
                inc_header = [
                    Paragraph("Tipo", table_cell_style),
                    Paragraph("Tipo Monto", table_cell_style),
                    Paragraph("Valor", table_cell_style),
                    Paragraph("Aplicación", table_cell_style),
                    Paragraph("Inicio", table_cell_style),
                    Paragraph("Fin", table_cell_style),
                    Paragraph("Descripción", table_cell_style),
                    Paragraph("Cantidad", table_cell_style),
                ]
                
                inc_table_data = [inc_header]
                for inc in increases:
                    inc_table_data.append([
                        Paragraph(getattr(inc.increase_type, 'name', 'N/A'), table_cell_style),
                        Paragraph(inc.amount_type, table_cell_style),
                        Paragraph(f"{inc.amount_value}", table_cell_style),
                        Paragraph(inc.application_increase_type, table_cell_style),
                        Paragraph(ContractDocumentGenerator._format_date(inc.start_date_increase), table_cell_style),
                        Paragraph(ContractDocumentGenerator._format_date(inc.end_date_increase), table_cell_style),
                        Paragraph(inc.description or "N/A", table_cell_style),
                        Paragraph(f"{inc.amount}" if inc.amount else "N/A", table_cell_style),
                    ])
                
                inc_col_widths = [available_width * 0.15, available_width * 0.12, available_width * 0.1, available_width * 0.12,
                                 available_width * 0.1, available_width * 0.1, available_width * 0.16, available_width * 0.15]
                inc_table = Table(inc_table_data, colWidths=inc_col_widths, hAlign='LEFT', repeatRows=1)
                inc_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('LEFTPADDING', (0, 0), (-1, -1), 4),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                ]))
                story.append(inc_table)
            
            return story
        
        # Primera pasada: construir para contar páginas
        doc.build(create_story(doc.width), onFirstPage=on_page_first, onLaterPages=on_page_first)
        
        # Obtener el total de páginas
        total_pages = page_info[0]
        
        # Si tenemos páginas, reconstruir con el total conocido usando un nuevo buffer
        if total_pages > 0:
            # Crear nuevo buffer y documento para la segunda construcción
            buffer_final = BytesIO()
            doc_final = SimpleDocTemplate(
                buffer_final,
                pagesize=A4,
                leftMargin=left_margin,
                rightMargin=right_margin,
                topMargin=top_margin,
                bottomMargin=bottom_margin,
            )
            doc_final.build(create_story(doc_final.width), onFirstPage=on_page_final, onLaterPages=on_page_final)
            pdf_bytes = buffer_final.getvalue()
            buffer_final.close()
        else:
            # Si no hay páginas (no debería pasar), usar el buffer original
            pdf_bytes = buffer.getvalue()
        
        buffer.close()
        return pdf_bytes
    
    @staticmethod
    def generate_docx(contract, downloader_user=None, logo_path=None):
        """
        Genera un DOCX del contrato establecido.
        
        Args:
            contract: Instancia de EstablishedContract
            downloader_user: Usuario que descarga (str, dict u objeto)
            logo_path: Ruta al logo (opcional)
        
        Returns:
            bytes: Contenido del DOCX
        """
        doc = Document()
        
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        # Encabezado con logo y título
        header = sections[0].header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if logo_path:
            try:
                header_para.add_run().add_picture(logo_path, width=Inches(1.2))
            except Exception:
                pass
        
        # Título principal
        title = doc.add_heading(f'Contrato de Trabajo {contract.contract_code}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Fecha de generación
        gen_date_para = doc.add_paragraph()
        gen_date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gen_date = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S")
        gen_date_run = gen_date_para.add_run(f'Generado: {gen_date}')
        gen_date_run.font.size = Pt(9)
        gen_date_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()
        
        # Sección 1: Generalidades del contrato
        doc.add_heading('1. GENERALIDADES DEL CONTRATO', level=1)
        
        generalidades_data = [
            ["Código del contrato", contract.contract_code],
            ["Cargo", getattr(contract.id_employee_charge, 'name', 'N/A')],
            ["Descripción", contract.description or "N/A"],
            ["Tipo de contrato", getattr(contract.contract_type, 'name', 'N/A')],
            ["Fecha de inicio", ContractDocumentGenerator._format_date(contract.start_date)],
            ["Fecha de finalización", ContractDocumentGenerator._format_date(contract.end_date)],
            ["Frecuencia de pago", ContractDocumentGenerator._format_payment_frequency_details(contract)],
            ["Jornada laboral", getattr(contract.workday_type, 'name', 'N/A') if contract.workday_type else "N/A"],
            ["Modalidad de trabajo", getattr(contract.work_mode_type, 'name', 'N/A') if contract.work_mode_type else "N/A"],
            ["Estado", getattr(contract.established_contract_status, 'name', 'N/A')],
        ]
        
        gen_table = doc.add_table(rows=len(generalidades_data), cols=2)
        gen_table.style = 'Light Grid Accent 1'
        gen_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        for i, (label, value) in enumerate(generalidades_data):
            gen_table.rows[i].cells[0].text = label
            gen_table.rows[i].cells[1].text = str(value)
            gen_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Sección 2: Términos del contrato
        doc.add_heading('2. TÉRMINOS DEL CONTRATO', level=1)
        
        currency_symbol = getattr(contract.currency_type, 'symbol', '') or ''
        salary_display = f"{currency_symbol}{contract.salary_base:,.2f}"
        
        terminos_data = [
            ["Tipo de salario", contract.salary_type],
            ["Salario base", salary_display],
            ["Moneda", getattr(contract.currency_type, 'name', 'N/A')],
            ["Período de prueba (días)", str(contract.trial_period_days) if contract.trial_period_days else "N/A"],
            ["Días de vacaciones", str(contract.vacation_days)],
            ["Vacaciones acumulativas", "Sí" if contract.cumulative_vacation else "No"],
            ["Fecha inicio acumulación", ContractDocumentGenerator._format_date(contract.start_cumulative_vacation) if contract.cumulative_vacation else "N/A"],
            ["Frecuencia de vacaciones (días)", str(contract.vacation_frequency_days) if contract.vacation_frequency_days else "N/A"],
            ["Días máximos de incapacidad", str(contract.maximum_disability_days)],
            ["Horas extras", f"{contract.overtime}" if contract.overtime else "0"],
            ["Período horas extras", contract.overtime_period or "N/A"],
            ["Período de notificación (días)", str(contract.notice_period_days) if contract.notice_period_days else "N/A"],
        ]
        
        if contract.minimum_hours:
            terminos_data.insert(3, ["Horas mínimas", str(contract.minimum_hours)])
        
        term_table = doc.add_table(rows=len(terminos_data), cols=2)
        term_table.style = 'Light Grid Accent 1'
        term_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        for i, (label, value) in enumerate(terminos_data):
            term_table.rows[i].cells[0].text = label
            term_table.rows[i].cells[1].text = str(value)
            term_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Sección 3: Deducciones
        deductions = list(contract.established_deductions.all())
        if deductions:
            doc.add_heading('3. DEDUCCIONES', level=1)
            
            ded_table = doc.add_table(rows=len(deductions) + 1, cols=8)
            ded_table.style = 'Light Grid Accent 1'
            
            # Encabezados
            headers = ["Tipo", "Tipo Monto", "Valor", "Aplicación", "Inicio", "Fin", "Descripción", "Cantidad"]
            for i, header in enumerate(headers):
                cell = ded_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Datos
            for idx, ded in enumerate(deductions, start=1):
                ded_table.rows[idx].cells[0].text = getattr(ded.deduction_type, 'name', 'N/A')
                ded_table.rows[idx].cells[1].text = ded.amount_type
                ded_table.rows[idx].cells[2].text = f"{ded.amount_value}"
                ded_table.rows[idx].cells[3].text = ded.application_deduction_type
                ded_table.rows[idx].cells[4].text = ContractDocumentGenerator._format_date(ded.start_date_deduction)
                ded_table.rows[idx].cells[5].text = ContractDocumentGenerator._format_date(ded.end_date_deductions)
                ded_table.rows[idx].cells[6].text = ded.description or "N/A"
                ded_table.rows[idx].cells[7].text = f"{ded.amount}" if ded.amount else "N/A"
            
            doc.add_paragraph()
        
        # Sección 4: Incrementos
        increases = list(contract.established_increases.all())
        if increases:
            doc.add_heading('4. INCREMENTOS', level=1)
            
            inc_table = doc.add_table(rows=len(increases) + 1, cols=8)
            inc_table.style = 'Light Grid Accent 1'
            
            # Encabezados
            headers = ["Tipo", "Tipo Monto", "Valor", "Aplicación", "Inicio", "Fin", "Descripción", "Cantidad"]
            for i, header in enumerate(headers):
                cell = inc_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Datos
            for idx, inc in enumerate(increases, start=1):
                inc_table.rows[idx].cells[0].text = getattr(inc.increase_type, 'name', 'N/A')
                inc_table.rows[idx].cells[1].text = inc.amount_type
                inc_table.rows[idx].cells[2].text = f"{inc.amount_value}"
                inc_table.rows[idx].cells[3].text = inc.application_increase_type
                inc_table.rows[idx].cells[4].text = ContractDocumentGenerator._format_date(inc.start_date_increase)
                inc_table.rows[idx].cells[5].text = ContractDocumentGenerator._format_date(inc.end_date_increase)
                inc_table.rows[idx].cells[6].text = inc.description or "N/A"
                inc_table.rows[idx].cells[7].text = f"{inc.amount}" if inc.amount else "N/A"
        
        # Pie de página
        footer = sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        user_name = ContractDocumentGenerator._format_user_name(downloader_user)
        
        # Agregar texto del usuario y paginación
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            # Texto del usuario
            user_run = footer_para.add_run(f'Descargado por: {user_name}')
            user_run.font.size = Pt(8)
            user_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Separador
            sep_run = footer_para.add_run(' | Página ')
            sep_run.font.size = Pt(8)
            sep_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Campo PAGE (página actual)
            page_run = footer_para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = ' PAGE '
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            page_run._element.append(fldChar1)
            page_run._element.append(instrText)
            page_run._element.append(fldChar2)
            page_run.font.size = Pt(8)
            page_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Separador para total
            sep2_run = footer_para.add_run(' / ')
            sep2_run.font.size = Pt(8)
            sep2_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Campo NUMPAGES (total de páginas)
            total_run = footer_para.add_run()
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'begin')
            
            instrText2 = OxmlElement('w:instrText')
            instrText2.set(qn('xml:space'), 'preserve')
            instrText2.text = ' NUMPAGES '
            
            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')
            
            total_run._element.append(fldChar3)
            total_run._element.append(instrText2)
            total_run._element.append(fldChar4)
            total_run.font.size = Pt(8)
            total_run.font.color.rgb = RGBColor(128, 128, 128)
            
        except Exception as e:
            # Si falla, mostrar solo el usuario
            logger.warning(f"Error al agregar paginación en DOCX: {str(e)}")
            footer_para.clear()
            user_run = footer_para.add_run(f'Descargado por: {user_name}')
            user_run.font.size = Pt(8)
            user_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Guardar en bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.getvalue()
        buffer.close()
        
        return docx_bytes

